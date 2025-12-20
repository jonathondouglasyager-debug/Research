"""
DOCUMENT ANALYSIS AGENT - Process PDFs, Excel, CSV, Word documents
Extracts text, tables, entities, and insights from various document formats
"""

import sys
from pathlib import Path
from typing import Dict

# Add parent directory to path
sys.path.append(str(Path(__file__).parent.parent))

from research_agent import ResearchAgent


class DocumentAnalysisAgent(ResearchAgent):
    """Agent specialized in analyzing documents (PDF, Excel, CSV, DOCX)"""

    def __init__(self, agent_id: str, question: str, investigation: str,
                 document_path: str = None, research_dir: str = None):
        """
        Initialize document analysis agent

        Args:
            agent_id: Unique agent identifier
            question: Research question to answer from document
            investigation: Parent investigation name
            document_path: Path to document to analyze
            research_dir: Research directory path
        """
        super().__init__(agent_id, question, investigation, research_dir)

        self.document_path = Path(document_path) if document_path else None
        self.agent_type = "document_analysis"

        # Update findings structure
        self.findings['document'] = {
            'path': str(self.document_path) if self.document_path else None,
            'type': self._detect_document_type(),
            'size': None,
            'pages': None
        }

    def _detect_document_type(self) -> str:
        """Detect document type from extension"""
        if not self.document_path or not self.document_path.exists():
            return 'unknown'

        extension = self.document_path.suffix.lower()

        type_mapping = {
            '.pdf': 'pdf',
            '.xlsx': 'excel',
            '.xls': 'excel',
            '.csv': 'csv',
            '.docx': 'word',
            '.doc': 'word',
            '.txt': 'text',
            '.md': 'markdown'
        }

        return type_mapping.get(extension, 'unknown')

    def execute_research(self) -> Dict:
        """
        Execute document analysis

        Returns:
            Dict with content, tables, and metadata
        """
        if not self.document_path or not self.document_path.exists():
            return self._create_error_result(f"Document not found: {self.document_path}")

        doc_type = self.findings['document']['type']

        # Route to appropriate handler
        if doc_type == 'pdf':
            return self._analyze_pdf()
        elif doc_type == 'excel':
            return self._analyze_excel()
        elif doc_type == 'csv':
            return self._analyze_csv()
        elif doc_type == 'word':
            return self._analyze_word()
        elif doc_type in ['text', 'markdown']:
            return self._analyze_text()
        else:
            return self._create_error_result(f"Unsupported document type: {doc_type}")

    def _analyze_pdf(self) -> Dict:
        """Analyze PDF document"""
        try:
            import PyPDF2

            with open(self.document_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                # Extract text from all pages
                full_text = ""
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    full_text += page.extract_text() + "\n\n"

                # Update document metadata
                self.findings['document']['pages'] = num_pages
                self.findings['document']['size'] = self.document_path.stat().st_size

                return {
                    'content': full_text,
                    'pages': num_pages,
                    'format': 'pdf',
                    'sources': [{
                        'title': self.document_path.name,
                        'url': f'file:///{str(self.document_path)}',
                        'type': 'pdf_document'
                    }]
                }

        except ImportError:
            return self._create_fallback_result("PyPDF2 not available - install with: pip install PyPDF2")
        except Exception as e:
            return self._create_error_result(f"PDF analysis error: {str(e)}")

    def _analyze_excel(self) -> Dict:
        """Analyze Excel document"""
        try:
            import pandas as pd

            # Read all sheets
            excel_file = pd.ExcelFile(self.document_path)
            sheet_names = excel_file.sheet_names

            content = f"Excel Document: {self.document_path.name}\n"
            content += f"Number of sheets: {len(sheet_names)}\n\n"

            tables = []

            for sheet_name in sheet_names:
                df = pd.read_excel(self.document_path, sheet_name=sheet_name)

                content += f"=== Sheet: {sheet_name} ===\n"
                content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                content += f"Columns: {', '.join(df.columns)}\n\n"

                # Sample data
                content += "Sample data (first 5 rows):\n"
                content += df.head().to_string() + "\n\n"

                # Basic statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    content += "Statistics:\n"
                    content += df[numeric_cols].describe().to_string() + "\n\n"

                # Store table data
                tables.append({
                    'sheet_name': sheet_name,
                    'rows': len(df),
                    'columns': list(df.columns),
                    'data': df.to_dict('records')[:100]  # First 100 rows
                })

            self.findings['document']['size'] = self.document_path.stat().st_size
            self.findings['tables'] = tables

            return {
                'content': content,
                'tables': tables,
                'format': 'excel',
                'sources': [{
                    'title': self.document_path.name,
                    'url': f'file:///{str(self.document_path)}',
                    'type': 'excel_document'
                }]
            }

        except ImportError:
            return self._create_fallback_result("pandas/openpyxl not available - install with: pip install pandas openpyxl")
        except Exception as e:
            return self._create_error_result(f"Excel analysis error: {str(e)}")

    def _analyze_csv(self) -> Dict:
        """Analyze CSV document"""
        try:
            import pandas as pd

            df = pd.read_csv(self.document_path)

            content = f"CSV Document: {self.document_path.name}\n"
            content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
            content += f"Columns: {', '.join(df.columns)}\n\n"

            # Sample data
            content += "Sample data (first 10 rows):\n"
            content += df.head(10).to_string() + "\n\n"

            # Basic statistics
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content += "Statistics:\n"
                content += df[numeric_cols].describe().to_string() + "\n\n"

            # Value counts for categorical columns
            categorical_cols = df.select_dtypes(include=['object']).columns
            for col in categorical_cols[:5]:  # First 5 categorical columns
                content += f"\nValue counts for '{col}':\n"
                content += df[col].value_counts().head(10).to_string() + "\n"

            self.findings['document']['size'] = self.document_path.stat().st_size
            self.findings['tables'] = [{
                'rows': len(df),
                'columns': list(df.columns),
                'data': df.to_dict('records')[:100]
            }]

            return {
                'content': content,
                'format': 'csv',
                'sources': [{
                    'title': self.document_path.name,
                    'url': f'file:///{str(self.document_path)}',
                    'type': 'csv_document'
                }]
            }

        except ImportError:
            return self._create_fallback_result("pandas not available - install with: pip install pandas")
        except Exception as e:
            return self._create_error_result(f"CSV analysis error: {str(e)}")

    def _analyze_word(self) -> Dict:
        """Analyze Word document"""
        try:
            from docx import Document

            doc = Document(self.document_path)

            # Extract all paragraphs
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            content = f"Word Document: {self.document_path.name}\n\n"
            content += full_text

            if tables:
                content += f"\n\nFound {len(tables)} table(s) in document\n"

            self.findings['document']['size'] = self.document_path.stat().st_size
            self.findings['tables'] = tables

            return {
                'content': content,
                'tables': tables,
                'format': 'word',
                'sources': [{
                    'title': self.document_path.name,
                    'url': f'file:///{str(self.document_path)}',
                    'type': 'word_document'
                }]
            }

        except ImportError:
            return self._create_fallback_result("python-docx not available - install with: pip install python-docx")
        except Exception as e:
            return self._create_error_result(f"Word analysis error: {str(e)}")

    def _analyze_text(self) -> Dict:
        """Analyze plain text or markdown document"""
        try:
            with open(self.document_path, 'r', encoding='utf-8') as f:
                content = f.read()

            self.findings['document']['size'] = self.document_path.stat().st_size

            return {
                'content': content,
                'format': 'text',
                'sources': [{
                    'title': self.document_path.name,
                    'url': f'file:///{str(self.document_path)}',
                    'type': 'text_document'
                }]
            }

        except Exception as e:
            return self._create_error_result(f"Text analysis error: {str(e)}")

    def _create_error_result(self, error_message: str) -> Dict:
        """Create error result"""
        return {
            'content': f"ERROR: {error_message}",
            'error': error_message,
            'format': 'error',
            'sources': []
        }

    def _create_fallback_result(self, message: str) -> Dict:
        """Create fallback result when libraries not available"""
        content = f"{message}\n\n"
        content += f"Document: {self.document_path.name}\n"
        content += f"Type: {self.findings['document']['type']}\n"
        content += f"Question: {self.question}\n\n"
        content += "FALLBACK MODE: Required libraries not installed.\n"
        content += "Document analysis placeholder generated.\n"

        return {
            'content': content,
            'fallback': True,
            'format': 'fallback',
            'sources': [{
                'title': self.document_path.name,
                'url': f'file:///{str(self.document_path)}',
                'type': 'document'
            }]
        }


def main():
    """Test document analysis agent"""
    import sys

    if len(sys.argv) < 4:
        print("Usage: python document_analysis_agent.py <agent_id> <document_path> <investigation>")
        print("Example: python document_analysis_agent.py doc_001 data.csv COVID_PCR_Truth_Investigation")
        return

    agent_id = sys.argv[1]
    document_path = sys.argv[2]
    investigation = sys.argv[3]
    question = f"Analyze document: {Path(document_path).name}"

    agent = DocumentAnalysisAgent(
        agent_id=agent_id,
        question=question,
        investigation=investigation,
        document_path=document_path
    )

    print(f"[DOC AGENT] Analyzing: {document_path}")
    result = agent.execute()

    print(f"\n[DOC AGENT] Analysis complete")
    print(f"Entities found: {len(result['entities'])}")
    print(f"Timeline events: {len(result['timeline_events'])}")
    print(f"Glossary terms: {len(result['glossary_terms'])}")


if __name__ == '__main__':
    main()
